"""황백화 AI 조기경보 — 학습 데이터셋 카탈로그 docx 생성.

실행:
    uv run python scripts/gen_dataset_docx.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
from pathlib import Path

OUT_PATH = Path("/mnt/c/Users/Administrator/Desktop/어텐션/docs/ML-데이터/dataset_catalog.docx")

NAVY    = RGBColor(0x1B, 0x36, 0x5D)
SUBNAVY = RGBColor(0x3A, 0x60, 0x73)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_L  = RGBColor(0xF2, 0xF4, 0xF7)

# 배경색용 hex 문자열 (set_cell_bg에 직접 전달)
BG_NAVY  = "1B365D"
BG_WHITE = "FFFFFF"
BG_GRAY  = "F2F4F7"


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(15)
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = SUBNAVY
        p.paragraph_format.space_before = Pt(10)
    else:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


def body(doc, text, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = italic
    p.paragraph_format.space_after = Pt(4)


def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"
    # 헤더
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_bg(cell, BG_NAVY)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 데이터
    for ri, row in enumerate(rows):
        bg = BG_GRAY if ri % 2 == 0 else BG_WHITE
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            set_cell_bg(cell, bg)
            cell.paragraphs[0].add_run(str(val)).font.size = Pt(9)
    # 열 너비
    if col_widths:
        for row in t.rows:
            for ci, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[ci])
    doc.add_paragraph()


def build():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)
    doc.styles["Normal"].font.name = "맑은 고딕"
    doc.styles["Normal"].font.size = Pt(10)

    # ── 표지 ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("황백화 AI 조기경보 시스템\n학습 데이터셋 카탈로그")
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("ATTENTIONPLZ — Ocean Tensor Cube 데이터 명세서")
    r2.font.size = Pt(12); r2.font.color.rgb = SUBNAVY

    doc.add_paragraph()
    pm = doc.add_paragraph()
    pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pm.add_run(
        f"작성일: {datetime.date.today()}  |  작성: 클또리(Claude Code)  |  구성안: 젬또리(Antigravity CLI)"
    ).font.size = Pt(9)
    doc.add_page_break()

    # ── 1. 개요 ───────────────────────────────────────────────────────────
    heading(doc, "1. 개요")
    body(doc,
        "본 문서는 서해 양식 김 황백화 AI 조기경보 모델 학습에 사용되는 Ocean Tensor Cube "
        "데이터셋의 버전별 명세, 채널 구성, 라벨링 정책을 기록한다. 모든 큐브는 "
        "(T, H, W, C) 형태의 4차원 텐서로, H100 서버 ~/cheolyoung/output/ 경로에 "
        "Zarr 형식으로 보관된다."
    )

    # ── 2. 데이터 소스 ────────────────────────────────────────────────────
    heading(doc, "2. 데이터 파이프라인 및 소스")
    table(doc,
        ["소스 기관", "정식 명칭", "제공 변수"],
        [
            ("NIFS",       "국립수산과학원",         "수온, DIN, DO, 염분 등"),
            ("KOEM",       "해양환경공단",           "DO, 탁도, pH, 영양염"),
            ("KMA",        "기상청",                "기온, 강수량, 풍속/풍향, 일사량"),
            ("K-water",    "한국수자원공사",         "하천방류량 (금강)"),
            ("GOCI-II",    "천리안 위성",            "클로로필-a, NIR 지수 (B865/B555)"),
            ("Sentinel-2", "ESA / Element84 STAC", "NDCI (B05-B04)/(B05+B04), 10m→2km"),
            ("KODC",       "한국해양자료센터",        "해류 월평균 (수동 다운로드)"),
        ],
        col_widths=[3.0, 5.0, 7.5]
    )
    body(doc,
        "수집 흐름: collect_only.py (로컬, ~23분) → SFTP → build_from_parquet.py (H100 IDW 보간, ~90분) → Zarr 저장",
        italic=True
    )
    body(doc,
        "⚠ KODC 해류: 웹 UI 수동 다운로드만 가능 → --local-nc 파라미터로 우회.\n"
        "⚠ Sentinel-2: Element84 STAC + AWS COG 스트리밍, cloud < 20% 자동 필터.",
        italic=True
    )
    doc.add_page_break()

    # ── 3. 버전별 명세 ────────────────────────────────────────────────────
    heading(doc, "3. 데이터셋 버전별 명세")
    table(doc,
        ["버전", "크기", "형태 (T×H×W×C)", "기간", "이벤트", "비고"],
        [
            ("cube_v3_parquet", "428 KB", "—",                    "2021-11 ~ 2026-01", "—",   "수집 원본 parquet. 큐브 빌드 소스"),
            ("cube_v1",         "7.7 GB", "(366, 128, 128, 32)",  "2022-11 ~ 2023-10", "1회", "초기 32채널. WQI 포함(이후 제거). 비압축 Zarr"),
            ("cube_v2",         "663 MB", "(486, 128, 128, 32)",  "2022-11 ~ 2024-02", "1회", "DIN 결측 forward-fill 이슈. Random split 시계열 누수"),
            ("cube_pre2021",    "663 MB", "(485, 128, 128, 33)",  "2016-11 ~ 2018-02", "1회", "2016-17 이벤트 포함. v3 병합 시 센서 이질성 노이즈"),
            ("cube_v3",         "2.1 GB", "(1553, 128, 128, 33)", "2021-11 ~ 2026-01", "3회", "현행 표준 데이터셋. Sentinel-2 NDCI 포함"),
            ("cube_final",      "2.7 GB", "(2038, 128, 128, 33)", "2016-11 ~ 2026-01", "4회", "pre2021+v3 병합. v6 학습 사용. 환경 이질성으로 성능 하락"),
        ],
        col_widths=[3.2, 1.6, 4.0, 3.5, 1.5, 5.5]
    )

    heading(doc, "황백화 이벤트 이력", 2)
    table(doc,
        ["회차", "기간", "포함 큐브"],
        [
            ("1회", "2016-11 ~ 2017-02", "cube_pre2021"),
            ("2회", "2022-01 ~ 2022-02", "cube_v3"),
            ("3회", "2023-11 ~ 2024-02", "cube_v3"),
            ("4회", "2025-11 ~ 2026-01", "cube_v3"),
        ],
        col_widths=[2.0, 5.0, 4.0]
    )
    doc.add_page_break()

    # ── 4. 채널 명세 ──────────────────────────────────────────────────────
    heading(doc, "4. 채널 명세 (33채널, cube_v3 기준)")
    body(doc, "학습 시 DIN 관측 마스크(din_mask) 채널이 자동 추가되어 실제 모델 입력은 34채널.")
    table(doc,
        ["#", "변수명", "단위", "소스", "분류"],
        [
            ("00", "sst",                "°C",         "NIFS + KOEM",              "Tier0 직접 원인"),
            ("01", "din",                "μmol/L",     "NIFS + KOEM",              "Tier0 직접 원인"),
            ("02", "dip",                "μmol/L",     "KOEM",                     "Tier0 직접 원인"),
            ("03", "sio2",               "μmol/L",     "KOEM",                     "Tier1 핵심 forcing"),
            ("04", "np_ratio",           "—",          "파생 (ch1÷ch2)",           "Tier1 핵심 forcing"),
            ("05", "salinity",           "psu",        "NIFS + KOEM",              "Tier1 핵심 forcing"),
            ("06", "precipitation",      "mm/day",     "KMA",                      "Tier1 핵심 forcing"),
            ("07", "discharge",          "m³/s",       "K-water (금강)",           "Tier1 핵심 forcing"),
            ("08", "dist_estuary",       "km",         "정적 계산",                "Tier1 핵심 forcing"),
            ("09", "par_proxy",          "μmol/m²/s", "KMA 일사량 근사",          "Tier1 핵심 forcing"),
            ("10", "chlorophyll_a",      "mg/m³",      "GOCI-II + KOEM",           "Tier2 보조 환경장"),
            ("11", "dissolved_oxygen",   "mg/L",       "NIFS + KOEM",              "Tier2 보조 환경장"),
            ("12", "current_u",          "m/s",        "KODC 월평균",              "Tier2 보조 환경장"),
            ("13", "current_v",          "m/s",        "KODC 월평균",              "Tier2 보조 환경장"),
            ("14", "water_depth",        "m",          "정적 계산",                "Tier2 보조 환경장"),
            ("15", "sst_anomaly",        "°C",         "파생 (ch0 - 30일 평균)",   "Tier2 보조 환경장"),
            ("16", "sst_7d_avg",         "°C",         "파생",                     "Tier2 보조 환경장"),
            ("17", "days_since_rain",    "day",        "파생 (ch6)",               "Tier2 보조 환경장"),
            ("18", "turbidity",          "NTU",        "KOEM",                     "Tier2 보조 환경장"),
            ("19", "wind_speed",         "m/s",        "KMA",                      "Tier3 운영/기상 보조"),
            ("20", "wind_dir_sin",       "—",          "KMA",                      "Tier3 운영/기상 보조"),
            ("21", "wind_dir_cos",       "—",          "KMA",                      "Tier3 운영/기상 보조"),
            ("22", "air_temp",           "°C",         "KMA",                      "Tier3 운영/기상 보조"),
            ("23", "ph",                 "—",          "KOEM",                     "Tier3 운영/기상 보조"),
            ("24", "no3_nitrogen",       "μmol/L",     "NIFS + KOEM",              "Tier3 운영/기상 보조"),
            ("25", "nh4_nitrogen",       "μmol/L",     "NIFS + KOEM",              "Tier3 운영/기상 보조"),
            ("26", "tn_proxy",           "μmol/L",     "파생 (no3+nh4+유기질소)",  "Tier3 운영/기상 보조"),
            ("27", "sst_gradient",       "°C/km",      "파생 (ch0 공간 미분)",     "Tier4 선택/파생"),
            ("28", "salinity_3d_change", "psu",        "파생 (ch5 3일 변화량)",    "Tier4 선택/파생"),
            ("29", "exposure_time",      "hr",         "파생 (조위 + 수심)",       "Tier4 선택/파생"),
            ("30", "growth_stage",       "—",          "파생 (작기 월기반 proxy)", "Tier4 선택/파생"),
            ("31", "nir_idx",            "—",          "GOCI-II B865/B555",        "Tier4 선택/파생"),
            ("32", "sentinel_ndci",      "—",          "Sentinel-2 B05/B04 NDCI",  "Tier4 선택/파생"),
            ("+1", "din_mask",           "—",          "파생 (DIN 실측일 마스크)", "학습 시 자동 추가"),
        ],
        col_widths=[0.8, 4.2, 2.2, 4.3, 3.8]
    )
    doc.add_page_break()

    # ── 5. 라벨링 정책 ────────────────────────────────────────────────────
    heading(doc, "5. 라벨링 및 학습/검증 분할 정책")
    heading(doc, "황백화 단계 정의", 2)
    table(doc,
        ["라벨", "이름", "기준"],
        [
            ("0",  "정상",   "황백화 비발생"),
            ("1",  "초기",   "이벤트 시작 후 0~4주"),
            ("2",  "경계",   "이벤트 시작 후 4~10주"),
            ("3",  "진행",   "이벤트 시작 후 10주 이상"),
            ("-1", "ignore", "비수확기 및 이벤트 버퍼 구간"),
        ],
        col_widths=[1.5, 2.5, 9.0]
    )

    heading(doc, "Stratified Temporal Split (80:20)", 2)
    body(doc,
        "• 라벨별로 앞 80% → train, 뒤 20% → val\n"
        "• 목적: 시계열 데이터 누수(Data Leakage) 방지 + 이벤트 클래스 val 포함 보장"
    )
    heading(doc, "cube_v3 기준 분포 (t_in=24일, stride=6일)", 3)
    table(doc,
        ["클래스", "전체 t_index", "train", "val"],
        [
            ("정상(0)", "210", "168", "42"),
            ("초기(1)", "14",  "11",  "3"),
            ("경계(2)", "20",  "16",  "4"),
            ("진행(3)", "11",  "8",   "3"),
        ],
        col_widths=[3.0, 3.5, 3.0, 3.0]
    )

    # ── 6. 서버 위치 ──────────────────────────────────────────────────────
    heading(doc, "6. 서버 보관 위치 (H100)")
    table(doc,
        ["경로", "설명"],
        [
            ("~/cheolyoung/output/cube_v3/",        "현행 표준 데이터셋 (Zarr)"),
            ("~/cheolyoung/output/cube_final/",      "pre2021 병합본 (Zarr)"),
            ("~/cheolyoung/output/cube_pre2021/",    "구버전 과거 이벤트 (Zarr)"),
            ("~/cheolyoung/output/cube_v3_parquet/", "수집 원본 parquet"),
        ],
        col_widths=[8.0, 7.5]
    )
    body(doc, "접속: ssh tta@123.41.22.216  |  인증 정보: seaweed-hwangbaek/.env.h100", italic=True)

    # ── 푸터 ─────────────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        f"ATTENTIONPLZ 내부 문서  |  {datetime.date.today()}  |  "
        "클또리(Claude Code) 작성 · 젬또리(Antigravity CLI) 구성안 · 챗또리(Codex) 리뷰"
    )
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"✅ 저장 완료: {OUT_PATH}")


if __name__ == "__main__":
    build()
