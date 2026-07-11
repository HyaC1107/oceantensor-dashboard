"""황백화 AI — ML 데이터셋 인수인계 문서 생성.

실행:
    uv run python scripts/gen_handover_docx.py
"""
import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_PATH = Path(
    "/mnt/c/Users/Administrator/Desktop/어텐션/docs/ML-데이터/데이터셋_인수인계.docx"
)

NAVY    = RGBColor(0x1B, 0x36, 0x5D)
SUBNAVY = RGBColor(0x3A, 0x60, 0x73)
RED     = RGBColor(0xC0, 0x39, 0x2B)
GREEN   = RGBColor(0x1A, 0x7A, 0x4A)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_L  = RGBColor(0xF2, 0xF4, 0xF7)

BG_NAVY   = "1B365D"
BG_RED    = "FDECEA"
BG_GREEN  = "EAF7EE"
BG_YELLOW = "FFF9E6"
BG_WHITE  = "FFFFFF"
BG_GRAY   = "F2F4F7"


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(5)


def h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = SUBNAVY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)


def h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)


def body(doc, text, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(3)


def code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_after  = Pt(4)
    # 배경 음영
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "EEF2F7")
    pPr.append(shd)


def tbl(doc, headers, rows, col_widths=None, row_colors=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_bg(cell, BG_NAVY)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        bg = (row_colors[ri] if row_colors else None) or (BG_GRAY if ri % 2 == 0 else BG_WHITE)
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            set_cell_bg(cell, bg)
            cell.paragraphs[0].add_run(str(val)).font.size = Pt(8.5)
    if col_widths:
        for row in t.rows:
            for ci, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[ci])
    doc.add_paragraph()


def callout(doc, text, bg="FFF9E6", text_color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.italic = True
    if text_color:
        run.font.color.rgb = text_color
    p.paragraph_format.left_indent  = Cm(0.6)
    p.paragraph_format.space_after  = Pt(5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  bg)
    pPr.append(shd)


def build():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)
    doc.styles["Normal"].font.name = "맑은 고딕"
    doc.styles["Normal"].font.size = Pt(10)

    # ── 표지 ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ATTENTIONPLZ\nML 데이터셋 인수인계 문서")
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Ocean Tensor Cube 구조 · 채널 명세 · 0값 원인 및 대처방안")
    r2.font.size = Pt(12); r2.font.color.rgb = SUBNAVY

    doc.add_paragraph()
    pm = doc.add_paragraph()
    pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pm.add_run(
        f"작성일: {datetime.date.today()}  |  작성: 클또리(Claude Code)"
    ).font.size = Pt(9)
    doc.add_page_break()

    # ── 1. 개요 ──
    h1(doc, "1. 프로젝트 개요")
    body(doc,
        "서해 양식 김의 황백화(white rot) 조기경보 AI 시스템. "
        "공공 API 5종과 위성 데이터 2종을 수집해 IDW 공간보간으로 "
        "128×128 격자의 4D 텐서(Ocean Tensor Cube)를 구성하고, "
        "ST-MMT(Spatio-Temporal Multi-Modal Transformer)로 황백화 단계를 예측한다."
    )
    tbl(doc,
        ["항목", "내용"],
        [
            ("연구 영역 (ROI)", "위도 34.0~36.8°N / 경도 125.3~128.0°E (서해 김 양식권)"),
            ("격자 해상도",     "128×128 (약 2.2km/격자)"),
            ("채널 수",         "33채널 (학습 시 din_mask 추가 → 34채널)"),
            ("라벨",            "0=정상 / 1=초기 / 2=경계 / 3=진행 / -1=ignore"),
            ("모델",            "ST-MMT (Spatio-Temporal Multi-Modal Transformer)"),
            ("현행 표준 큐브",  "cube_v3 — (1553, 128, 128, 33), 2021-11 ~ 2026-01"),
        ],
        col_widths=[3.8, 11.7]
    )

    # ── 2. 파이프라인 ──
    h1(doc, "2. 데이터 파이프라인")
    body(doc, "전체 흐름: 로컬 수집(~23분) → SFTP 전송 → H100 보간/빌드(~90분) → Zarr 저장")
    h2(doc, "파이프라인 단계")
    tbl(doc,
        ["단계", "스크립트", "실행 위치", "소요 시간"],
        [
            ("① API 수집",    "collect_only.py",       "로컬 PC",    "~23분"),
            ("② SFTP 전송",   "scp 명령",               "로컬 PC",    "~5분"),
            ("③ IDW 보간",    "build_from_parquet.py", "H100 서버",  "~90분"),
            ("④ 학습",        "train_real.py",          "H100 서버",  "~15분/epoch"),
        ],
        col_widths=[3.0, 4.5, 4.0, 4.0]
    )
    h2(doc, "실행 명령")
    code_block(doc,
        "# 1. 로컬 수집\n"
        "PYTHONPATH=. python scripts/collect_only.py \\\n"
        "    --start-date 2021-11-01 --end-date 2026-01-31\n\n"
        "# 2. H100 전송\n"
        "scp -r output/checkpoints tta@123.41.22.216:/data/tta/shared/\n\n"
        "# 3. H100 보간 (H100 접속 후)\n"
        "PYTHONPATH=. python scripts/build_from_parquet.py \\\n"
        "    --checkpoint-dir /data/tta/shared/checkpoints"
    )
    doc.add_page_break()

    # ── 3. 데이터 소스 ──
    h1(doc, "3. 데이터 소스 (수집기별)")
    tbl(doc,
        ["수집기", "기관", "주요 변수", "갱신 주기", "비고"],
        [
            ("nifs_ml.py",    "NIFS 국립수산과학원",   "수온, 염분, DO, NO3, NH4",          "분기 1회", "femoSeaList API"),
            ("koem_ml.py",    "KOEM 해양환경공단",     "DO, 탁도, pH, DIN, DIP, SiO2, Chl-a", "월 1회",   "getOceansNemo2"),
            ("kma_ml.py",     "KMA 기상청",            "기온, 강수, 풍속/풍향, 일사량",      "일 1회",   "ASOS 11개 관측소"),
            ("kwater_ml.py",  "K-water",               "금강 방류량",                        "일 1회",   "하천수위자료"),
            ("kosc_ml.py",    "KOSC 국가해양위성센터", "Chl-a(OC3), NIR 지수",              "일 10회 합성", "GOCI-II OPeNDAP"),
            ("sentinel_ml.py","ESA / Element84 STAC",  "NDCI (B05-B04)/(B05+B04)",           "약 5일 1회", "AWS COG 스트리밍, cloud<20% 자동 필터"),
            ("nifs_ml.py\n(KODC)", "한국해양자료센터", "해류 월평균 u/v",                  "연 1회 갱신", "⚠️ 수동 다운로드 필수"),
        ],
        col_widths=[2.8, 3.8, 4.5, 2.3, 4.1]
    )
    callout(doc,
        "⚠️ KODC 해류: OPeNDAP 직접 접근 불가. 웹 UI 수동 다운로드 후 --local-nc 파라미터로 지정 필수.\n"
        "⚠️ GOCI-II: 구름 100% 날은 합성 스킵. 연속 흐린 날 forward-fill로 보완.",
        bg="FDECEA"
    )
    doc.add_page_break()

    # ── 4. 33채널 ──
    h1(doc, "4. 33채널 구조")
    body(doc, "학습 시 DIN 실측 마스크(din_mask)가 자동 추가되어 실제 모델 입력은 34채널.")

    h2(doc, "Tier0 — 직접 원인 (ch00~02)")
    tbl(doc,
        ["#", "변수명", "단위", "소스", "설명"],
        [
            ("00", "sst",  "°C",      "NIFS + KOEM", "표층 수온. 황백화 직접 유발 인자"),
            ("01", "din",  "μmol/L",  "NIFS + KOEM", "용존무기질소 (NO3+NH4). 분기 측정"),
            ("02", "dip",  "μmol/L",  "KOEM",        "용존무기인. 월 측정"),
        ],
        col_widths=[0.8, 3.8, 2.0, 4.0, 5.9]
    )

    h2(doc, "Tier1 — 핵심 forcing (ch03~09)")
    tbl(doc,
        ["#", "변수명", "단위", "소스", "설명"],
        [
            ("03", "sio2",          "μmol/L",     "KOEM",               "규산염"),
            ("04", "np_ratio",      "—",          "파생 (ch1÷ch2)",     "N:P 비율. dip=0이면 NaN"),
            ("05", "salinity",      "psu",        "NIFS + KOEM",        "염분"),
            ("06", "precipitation", "mm/day",     "KMA",                "강수량. 맑은 날 실제 0"),
            ("07", "discharge",     "m³/s",       "K-water (금강)",     "하천 방류량"),
            ("08", "dist_estuary",  "km",         "정적 계산",          "금강 하구까지 거리. 시간 불변"),
            ("09", "par_proxy",     "μmol/m²/s", "KMA 일사량 근사",    "일사량×26.6 변환"),
        ],
        col_widths=[0.8, 3.8, 2.0, 4.0, 5.9]
    )

    h2(doc, "Tier2 — 보조 환경장 (ch10~18)")
    tbl(doc,
        ["#", "변수명", "단위", "소스", "설명"],
        [
            ("10", "chlorophyll_a",    "mg/m³",  "GOCI-II + KOEM",    "클로로필-a"),
            ("11", "dissolved_oxygen", "mg/L",   "NIFS + KOEM",       "용존산소"),
            ("12", "current_u",        "m/s",    "KODC 월평균",       "동서류 ⚠️ 수동 다운로드"),
            ("13", "current_v",        "m/s",    "KODC 월평균",       "남북류 ⚠️ 수동 다운로드"),
            ("14", "water_depth",      "m",      "정적 계산",         "수심 근사 (5~100m 클리핑)"),
            ("15", "sst_anomaly",      "°C",     "파생 (ch0-30일평균)","수온 편차"),
            ("16", "sst_7d_avg",       "°C",     "파생",              "7일 평균 수온"),
            ("17", "days_since_rain",  "day",    "파생 (ch6)",        "강우 후 경과일. 비 온 당일=0"),
            ("18", "turbidity",        "NTU",    "KOEM",              "탁도"),
        ],
        col_widths=[0.8, 3.8, 2.0, 4.0, 5.9]
    )

    h2(doc, "Tier3 — 운영/기상 보조 (ch19~26)")
    tbl(doc,
        ["#", "변수명", "단위", "소스", "설명"],
        [
            ("19", "wind_speed",    "m/s",    "KMA", "풍속"),
            ("20", "wind_dir_sin",  "—",      "KMA", "풍향 sin. 북/남=0 (물리적 정상값)"),
            ("21", "wind_dir_cos",  "—",      "KMA", "풍향 cos. 동/서=0 (물리적 정상값)"),
            ("22", "air_temp",      "°C",     "KMA", "기온"),
            ("23", "ph",            "—",      "KOEM","pH"),
            ("24", "no3_nitrogen",  "μmol/L", "NIFS + KOEM", "NO3"),
            ("25", "nh4_nitrogen",  "μmol/L", "NIFS + KOEM", "NH4"),
            ("26", "tn_proxy",      "μmol/L", "파생","총질소 근사 = (NO3+NH4)×1.3"),
        ],
        col_widths=[0.8, 3.8, 2.0, 4.0, 5.9]
    )

    h2(doc, "Tier4 — 선택/파생 (ch27~32)")
    tbl(doc,
        ["#", "변수명", "단위", "소스", "설명"],
        [
            ("27", "sst_gradient",       "°C/km", "파생 (ch0 공간미분)",  "수온 공간 경사"),
            ("28", "salinity_3d_change",  "psu",   "파생 (ch5 3일변화)",   "염분 3일 변화. 첫 3일=0"),
            ("29", "exposure_time",       "hr",    "파생 (수심 기반)",     "공기 노출시간. 수심≥5m→0"),
            ("30", "growth_stage",        "—",     "파생 (월 proxy)",      "작기 경과일 (채묘=10월 기준)"),
            ("31", "nir_idx",             "—",     "GOCI-II B865/B555",   "NIR 지수. 황백화 시 상승"),
            ("32", "sentinel_ndci",       "—",     "Sentinel-2 B05/B04",  "NDCI. 고해상도 Chl-a proxy"),
        ],
        col_widths=[0.8, 3.8, 2.0, 4.0, 5.9]
    )
    doc.add_page_break()

    # ── 5. 데이터셋 버전 ──
    h1(doc, "5. 데이터셋 버전")
    tbl(doc,
        ["버전", "형태 (T×H×W×C)", "기간", "이벤트", "상태"],
        [
            ("cube_v1",      "(366, 128, 128, 32)",  "2022-11 ~ 2023-10", "1회", "초기버전. WQI 포함(제거됨)"),
            ("cube_v2",      "(486, 128, 128, 32)",  "2022-11 ~ 2024-02", "1회", "DIN 결측 이슈. Random split 누수"),
            ("cube_pre2021", "(485, 128, 128, 33)",  "2016-11 ~ 2018-02", "1회", "과거 이벤트 포함. v3 병합 시 노이즈"),
            ("cube_v3",      "(1553, 128, 128, 33)", "2021-11 ~ 2026-01", "3회", "✅ 현행 표준. 권장 사용"),
            ("cube_final",   "(2038, 128, 128, 33)", "2016-11 ~ 2026-01", "4회", "pre2021+v3 병합. 환경 이질성으로 성능 하락"),
        ],
        col_widths=[3.0, 4.2, 3.8, 1.8, 4.7]
    )

    h2(doc, "황백화 이벤트 이력")
    tbl(doc,
        ["회차", "기간", "주요 지역", "포함 큐브"],
        [
            ("1회", "2016-11 ~ 2017-02", "서천, 군산",         "cube_pre2021"),
            ("2회", "2022-01 ~ 2022-02", "해남, 서천",         "cube_v3"),
            ("3회", "2023-11 ~ 2024-02", "서천, 충남전역",     "cube_v3"),
            ("4회", "2025-11 ~ 2026-01", "고흥, 군산, 서천",   "cube_v3"),
        ],
        col_widths=[2.0, 4.5, 4.5, 4.5]
    )
    doc.add_page_break()

    # ── 6. 라벨링 ──
    h1(doc, "6. 라벨링 정책")
    tbl(doc,
        ["라벨", "이름", "기준", "비고"],
        [
            ("0",  "정상",   "황백화 비발생 수확기",         ""),
            ("1",  "초기",   "이벤트 시작 후 0~4주",         "11월 초, 신호 미약"),
            ("2",  "경계",   "이벤트 시작 후 4~10주",        "12월~1월 초"),
            ("3",  "진행",   "이벤트 시작 후 10주 이상",     "1월 중순~2월"),
            ("-1", "ignore", "비수확기·채묘기·이벤트 전후 버퍼", "학습 제외"),
        ],
        col_widths=[1.5, 2.0, 5.5, 6.5]
    )
    callout(doc,
        "이벤트 시작 14일 전 / 종료 21일 후는 ignore(-1) 처리 (전이 구간 오염 방지).\n"
        "비수확기(6~8월), 채묘기(9~10월)는 전부 ignore.",
        bg="FFF9E6"
    )

    h2(doc, "학습/검증 분할 — Stratified Temporal Split 80:20")
    body(doc,
        "• 라벨별로 앞 80% → train, 뒤 20% → val (시계열 데이터 누수 방지)\n"
        "• t_in=24일, stride=6일 기준 cube_v3 분포:"
    )
    tbl(doc,
        ["클래스", "전체 t_index", "train", "val"],
        [
            ("정상(0)", "210", "168", "42"),
            ("초기(1)", "14",  "11",  "3"),
            ("경계(2)", "20",  "16",  "4"),
            ("진행(3)", "11",  "8",   "3"),
        ],
        col_widths=[4.0, 4.0, 3.5, 4.0]
    )
    doc.add_page_break()

    # ── 7. 0값 분석 ── (핵심 섹션)
    h1(doc, "7. 샘플 데이터 0값 — 원인 및 대처방안")
    callout(doc,
        "가장 많이 받는 질문: \"왜 이 채널이 전부 0이에요?\"\n"
        "0값은 세 종류가 있다: ① 정상 0 (물리적 의미)  ② 결측 대체 0 (문제)  ③ 초기화 0 (버퍼 부족)",
        bg="EEF2F7"
    )

    h2(doc, "7-A. 정상 0 — 물리적으로 의미 있는 값 ✅")
    body(doc, "이 채널들의 0은 실제 현상을 반영하므로 수정 불필요.")
    tbl(doc,
        ["채널", "변수", "0이 되는 상황", "확인 방법"],
        [
            ("06", "precipitation",  "맑은 날 (강수 없음)",            "여름철 외 빈번. 정상"),
            ("07", "discharge",      "갈수기 금강 방류 최소",           "K-water 원데이터 확인"),
            ("17", "days_since_rain","비 온 당일 카운터 리셋",          "0→N 증가 패턴이면 정상"),
            ("20", "wind_dir_sin",   "풍향이 북(0°) 또는 남(180°)",    "cos 채널 교차 확인"),
            ("21", "wind_dir_cos",   "풍향이 동(90°) 또는 서(270°)",   "sin 채널 교차 확인"),
            ("29", "exposure_time",  "수심 ≥5m 격자 (대부분 외해)",    "정적 채널, 모든 날 동일"),
        ],
        col_widths=[1.5, 4.0, 5.5, 5.5],
        row_colors=[BG_GREEN]*6
    )

    h2(doc, "7-B. 결측 대체 0 — 데이터 누락 → 강제 0 채움 ⚠️")
    body(doc, "channel_builder.py 결측 처리 마지막 단계에서 발생하는 위험한 0.")
    code_block(doc,
        "# channel_builder.py — 결측 처리 마지막 단계\n"
        "global_mean = np.nanmean(ch)\n"
        "if np.isnan(global_mean):   # 채널 전체가 NaN이면\n"
        "    global_mean = 0.0        # → 0으로 강제 채움 ← 위험!\n"
        "cube[:,:,:,c] = np.where(np.isnan(ch), global_mean, ch)"
    )
    body(doc, "아래 채널들은 소스 데이터가 없으면 전체 NaN → 전체 0이 된다.")
    tbl(doc,
        ["채널", "변수", "0이 되는 원인", "대처방안"],
        [
            ("12", "current_u",
             "KODC .nc 파일 미로드 시 전체 NaN → 0",
             "--local-nc 파라미터로 KODC 파일 지정 필수"),
            ("13", "current_v",
             "위와 동일",
             "위와 동일"),
            ("31", "nir_idx",
             "GOCI-II 미수집 또는 전 기간 구름 100%",
             "kosc.parquet 생성 확인. 구름 많은 계절 허용 범위 내"),
            ("32", "sentinel_ndci",
             "--no-sentinel 옵션 사용 시 전체 NaN → 0",
             "--no-sentinel 옵션 절대 사용 금지"),
            ("09", "par_proxy",
             "KMA solar_radiation_mjm2 필드 미응답",
             "KMA API 할당량 초과 여부 확인 후 재수집"),
        ],
        col_widths=[1.5, 4.0, 5.5, 5.5],
        row_colors=[BG_RED]*5
    )

    h3(doc, "진단 스크립트")
    body(doc, "큐브 빌드 후 반드시 실행해서 0값 비율 확인:")
    code_block(doc,
        "import zarr, numpy as np\n\n"
        "z = zarr.open('output/cube_v3/', mode='r')\n"
        "data = z['data']  # (T, H, W, C)\n\n"
        "NAMES = ['sst','din','dip','sio2','np_ratio','salinity',\n"
        "         'precipitation','discharge','dist_estuary','par_proxy',\n"
        "         'chlorophyll_a','dissolved_oxygen','current_u','current_v',\n"
        "         'water_depth','sst_anomaly','sst_7d_avg','days_since_rain',\n"
        "         'turbidity','wind_speed','wind_dir_sin','wind_dir_cos',\n"
        "         'air_temp','ph','no3_nitrogen','nh4_nitrogen','tn_proxy',\n"
        "         'sst_gradient','salinity_3d_change','exposure_time',\n"
        "         'growth_stage','nir_idx','sentinel_ndci']\n\n"
        "print('=== 채널별 0값 비율 ===')\n"
        "for c, name in enumerate(NAMES):\n"
        "    chunk = data[:, :, :, c][:]\n"
        "    pct = (chunk == 0).mean() * 100\n"
        "    mark = '⚠️ ' if pct > 5 else '✅ '\n"
        "    print(f'{mark} ch{c:02d} {name:25s}: {pct:.1f}%')"
    )

    h2(doc, "7-C. 초기화 0 — 수집 시작 시점 버퍼 부족")
    body(doc, "시계열 파생 채널은 이전 값이 쌓여야 계산되므로 수집 초기 며칠은 0.")
    tbl(doc,
        ["채널", "변수", "0이 되는 기간", "원인 및 영향"],
        [
            ("15", "sst_anomaly",      "첫 1일",  "자기 자신 대비 편차=0. 학습 영향 미미"),
            ("28", "salinity_3d_change","첫 3일",  "3일 전 값 없어 계산 불가. stride=6일 사용 시 자연 회피"),
        ],
        col_widths=[1.5, 4.0, 2.5, 8.5],
        row_colors=[BG_YELLOW]*2
    )

    h2(doc, "7-D. 체크리스트 (큐브 빌드 전)")
    body(doc,
        "□  KODC .nc 파일 존재 확인 → build_from_parquet.py --local-nc <경로>\n"
        "□  kosc.parquet 파일 크기 > 0 KB 확인 (ch31)\n"
        "□  --no-sentinel 옵션 미사용 확인 (ch32)\n"
        "□  parquet 5종 모두 존재: nifs / koem / kma / kwater / sentinel\n"
        "□  큐브 빌드 완료 후 위 7-B 진단 스크립트 실행"
    )
    doc.add_page_break()

    # ── 8. 학습 이력 ──
    h1(doc, "8. 학습 이력 요약")
    tbl(doc,
        ["버전", "큐브", "주요 설정", "F1", "AUC", "Accuracy", "비고"],
        [
            ("v5 (합성)", "synth",     "d_model=128, 20ep",                        "0.694", "0.944", "0.765", "파이프라인 검증용"),
            ("v6",        "cube_final","d_model=128, lr=1e-4",                      "0.322", "0.667", "0.725", "pre2021 환경 이질성 성능 하락"),
            ("v7",        "cube_v3",   "d_model=256, lr=1e-4, class_weight[진행]=6","0.622", "0.736", "0.911", "진행 F1=0.00 미해결"),
        ],
        col_widths=[1.8, 2.8, 5.5, 1.3, 1.3, 2.3, 4.5]
    )
    callout(doc,
        "v7 핵심 한계: 진행(class 3) F1 = 0.00 — 196,608 픽셀 전부 경계(2)로 오분류.\n"
        "원인: class_weight[진행]=6.0 부족 + 진행 t_index 11개로 절대 샘플 부족.\n"
        "v8 방향: class_weight[진행] 15~20 + FocalLoss gamma 3~5 + 진행 구간 오버샘플링.",
        bg="FDECEA"
    )

    # ── 9. 서버 정보 ──
    h1(doc, "9. 서버 정보")
    tbl(doc,
        ["항목", "내용"],
        [
            ("접속",       "ssh tta@123.41.22.216"),
            ("인증",       ".env.h100 참조 (⚠️ 비밀번호 끝 마침표가 비밀번호 일부)"),
            ("개인 작업공간", "/data/tta/cheolyoung/"),
            ("팀 공유",    "/data/tta/shared/  (parquet, 체크포인트)"),
            ("표준 큐브",  "~/cheolyoung/output/cube_v3/"),
            ("학습 체크포인트", "~/cheolyoung/checkpoints/"),
        ],
        col_widths=[4.0, 12.5]
    )

    # ── 10. 자주 묻는 질문 ──
    h1(doc, "10. 자주 묻는 질문")
    faq = [
        ("KODC 해류 데이터는 어떻게 받나요?",
         "KODC 웹 UI에서 수동 다운로드만 가능 (OPeNDAP 직접 접근 차단).\n"
         ".nc 파일 받아서 build_from_parquet.py --local-nc <파일경로> 로 지정."),
        ("Sentinel-2 데이터 없는 날 처리는?",
         "약 5일 주기 재방문. 없는 날은 forward-fill(직전 관측값 유지).\n"
         "tol_days=5 설정으로 5일 이내 관측 재사용. 그 이후엔 global_mean 채움."),
        ("GOCI-II 구름 덮임이 심한 날은?",
         "맑은 픽셀 비율 < 0.3%이면 해당 시각 스킵. 하루 10장 합성 시 일부만 사용.\n"
         "전체 합성 실패 시 forward-fill로 직전 날 합성값 사용."),
        ("cube_v3에 이벤트가 몇 회 들어 있나요?",
         "3회: 2022-01~02 / 2023-11~2024-02 / 2025-11~2026-01.\n"
         "cube_pre2021에 1회 추가 (2016-11~2017-02)."),
        ("din_mask 채널은 뭔가요?",
         "DIN은 분기 측정이라 대부분 날이 보간값. din_mask=1이면 실측일, =0이면 보간.\n"
         "학습 시 34번째 채널로 자동 concat됨."),
    ]
    for q, a in faq:
        h3(doc, f"Q. {q}")
        body(doc, f"A. {a}")

    # ── 푸터 ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        f"ATTENTIONPLZ 내부 문서  |  {datetime.date.today()}  |  "
        "클또리(Claude Code) 작성"
    )
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"✅ 저장 완료: {OUT_PATH}")


if __name__ == "__main__":
    build()
